"""Docx 解析器"""
from loguru import logger
from typing import List, Dict, Any

from .base import BaseParser



class DocxParser(BaseParser):
    name = 'docx'

    def parse(self, file_path: str, **options) -> List[Dict[str, Any]]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return []

        try:
            doc = DocxDocument(file_path)
        except Exception:
            logger.exception('open docx fail')
            return []

        blocks: List[Dict[str, Any]] = []
        current_section = ''
        for para in doc.paragraphs:
            text = (para.text or '').strip()
            if not text:
                continue
            style = (para.style.name or '').lower() if para.style else ''
            if 'heading' in style or style.startswith('title'):
                current_section = text[:64]
                blocks.append({
                    'type': 'text', 'content': text,
                    'section_path': current_section,
                    'page_number': None,
                    'extra': {'style': style},
                })
            else:
                blocks.append({
                    'type': 'text', 'content': text,
                    'section_path': current_section,
                    'page_number': None,
                    'extra': {},
                })
        # 表格
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append(' | '.join(cell.text.strip() for cell in row.cells))
            content = '\n'.join(rows)
            if content:
                blocks.append({
                    'type': 'table', 'content': content,
                    'section_path': current_section,
                    'page_number': None,
                    'extra': {'table_index': i},
                })
        return blocks
